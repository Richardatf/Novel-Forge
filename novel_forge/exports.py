from __future__ import annotations

import re
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from ebooklib import epub


PLACEHOLDERS = re.compile(r"\b(?:TODO|TBD|FIXME|\[INSERT[^\]]*\])\b", re.IGNORECASE)


def preflight(chapters: list[dict]) -> list[str]:
    findings: list[str] = []
    if not chapters:
        return ["No chapters are available for export."]
    numbers = sorted(int(c["chapter_number"]) for c in chapters)
    expected = set(range(numbers[0], numbers[-1] + 1))
    missing = sorted(expected - set(numbers))
    if missing:
        findings.append("Missing chapter numbers: " + ", ".join(map(str, missing)))
    for chapter in chapters:
        label = f"Chapter {chapter['chapter_number']}"
        if not str(chapter.get("title", "")).strip(): findings.append(f"{label} has a blank heading.")
        content = str(chapter.get("content", ""))
        if not content.strip(): findings.append(f"{label} is empty.")
        if PLACEHOLDERS.search(content): findings.append(f"{label} contains an unresolved placeholder.")
    return findings


def create_ebook_docx(chapters: list[dict], destination: Path, title: str, author: str = "") -> Path:
    doc = Document()
    doc.core_properties.title = title
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"; normal.font.size = Pt(11)
    doc.add_heading(title, 0)
    if author: doc.add_paragraph(author)
    for index, chapter in enumerate(chapters):
        if index: doc.add_page_break()
        doc.add_heading(f"Chapter {chapter['chapter_number']}: {chapter['title']}", level=1)
        for paragraph in str(chapter["content"]).split("\n\n"):
            if paragraph.strip(): doc.add_paragraph(paragraph.strip())
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)
    return destination


def create_print_docx(chapters: list[dict], destination: Path, title: str, trim=(6.0, 9.0), gutter=0.25) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(trim[0]); section.page_height = Inches(trim[1])
    section.top_margin = Inches(.75); section.bottom_margin = Inches(.75)
    section.left_margin = Inches(.75 + gutter); section.right_margin = Inches(.75)
    section.gutter = Inches(gutter)
    doc.add_heading(title, 0)
    for chapter in chapters:
        doc.add_page_break()
        doc.add_heading(f"Chapter {chapter['chapter_number']}\n{chapter['title']}", level=1)
        for paragraph in str(chapter["content"]).split("\n\n"):
            if paragraph.strip(): doc.add_paragraph(paragraph.strip())
    destination.parent.mkdir(parents=True, exist_ok=True); doc.save(destination)
    return destination


def create_epub(chapters: list[dict], destination: Path, title: str, author: str = "") -> Path:
    book = epub.EpubBook(); book.set_identifier(f"novel-forge-{title}"); book.set_title(title); book.set_language("en")
    if author: book.add_author(author)
    items = []
    for chapter in chapters:
        item = epub.EpubHtml(title=f"Chapter {chapter['chapter_number']}: {chapter['title']}", file_name=f"chapter-{chapter['chapter_number']}.xhtml", lang="en")
        paras = "".join(f"<p>{escape(p.strip())}</p>" for p in str(chapter["content"]).split("\n\n") if p.strip())
        item.content = f"<h1>Chapter {chapter['chapter_number']}: {chapter['title']}</h1>{paras}"
        book.add_item(item); items.append(item)
    book.toc = tuple(items); book.spine = ["nav", *items]
    book.add_item(epub.EpubNcx()); book.add_item(epub.EpubNav())
    destination.parent.mkdir(parents=True, exist_ok=True); epub.write_epub(destination, book)
    return destination
